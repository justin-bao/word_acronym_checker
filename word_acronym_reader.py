#!/usr/bin/env python3

# Program for reading all the acronyms in a .docx file
import re
from docx import Document
from docx.shared import Inches

def get_all_acronyms(path):
    """
    Find every (possible) acronym in the document
    (i.e. in the form of 'An Example Acronym (AEA)')
    Looks for any alphanumeric text surrounded by parentheses
    """

    document = Document(path)
    document.save(path)

    acronyms = []
    pattern = '[A-Z]\w*[A-Z]'

    for para in document.paragraphs:
        for acronym in re.findall(pattern, para.text):
            if acronym not in acronyms:
                acronyms.append(acronym)

    return acronyms


def get_explained_acronyms(path):
    """
    Find every acronym that has been explained
    (i.e. in the form of 'An Example Acronym (AEA)')
    """

    document = Document(path)
    document.save(path)

    acronymTuples = []
    pattern = '\([A-Z]\w*[A-Z]\)'

    for para in document.paragraphs:
        for acronym in re.findall(pattern, para.text):
            # backtrack until finding the number of words corresponding with the
            # number of capital letters in the acronym
            acronymIndex = para.text.index(acronym)
            numOfWords = sum(1 for c in acronym if c.isupper())


            explanation = ""
            acronymIndex -= 1

            # go until reaching the start of the string or until all words are found
            while acronymIndex >= 0 and numOfWords > 0:
                if para.text[acronymIndex].isupper():
                    numOfWords -= 1
                explanation = para.text[acronymIndex] + explanation
                acronymIndex -= 1

            # remove parentheses from the acronym and leading/trailing whitespace
            acronymTuples.append((acronym[1:-1], explanation.strip()))

    return dict(acronymTuples)


if __name__=="__main__":
    import sys
    print("get_all_acronyms:")
    print(get_all_acronyms(sys.argv[1]))
    print

    print("get_explained_acronyms:")
    print(get_explained_acronyms(sys.argv[1]))
    print
